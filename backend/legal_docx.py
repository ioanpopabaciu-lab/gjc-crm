"""
GJC Legal DOCX Generator
Generează documente Word (.docx) cu antet GJC, format juridic românesc.
"""

import io
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

UPLOADS_DIR = Path(__file__).parent / "uploads" / "legal"
UPLOADS_DIR.mkdir(parents=True, exist_ok=True)

# Date GJC pentru antet
GJC_NAME    = "GLOBAL JOBS CONSULTING SRL"
GJC_CUI     = "44678741"
GJC_ADDRESS = "Oradea, Bihor, România"
GJC_EMAIL   = "contact@gjc.ro"
GJC_PHONE   = ""
GJC_REG     = "J05/xxx/2023"


def _set_font(run, name: str = "Times New Roman", size: int = 12, bold: bool = False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold


def _add_paragraph(doc: "Document", text: str, align=WD_ALIGN_PARAGRAPH.LEFT,
                   bold: bool = False, size: int = 12, space_before: float = 0,
                   space_after: float = 6) -> "Paragraph":
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold)
    return p


def _add_horizontal_line(doc: "Document"):
    """Adaugă o linie orizontală."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_docx(
    title: str,
    body_text: str,
    template_id: str = "",
    variables: Optional[dict] = None,
    doc_id: str = "",
    emitent: str = "GJC",           # "GJC" sau "candidat"
    candidat_name: str = "",
) -> str:
    """
    Generează fișier .docx și returnează calea relativă.
    body_text: textul generat de Claude (cu citări inline).
    emitent: "GJC" → antet GJC cu ștampilă/semnătură; "candidat" → date candidat
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx nu este instalat. Adaugă python-docx în requirements_prod.txt")

    doc = Document()

    # ── Margini document ──────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

    # ── ANTET ─────────────────────────────────────────────────────────────────
    if emitent == "GJC":
        # Logo placeholder (text dacă nu există imagine)
        logo_path = Path(__file__).parent / "uploads" / "gjc_logo.png"
        if logo_path.exists():
            try:
                doc.add_picture(str(logo_path), width=Cm(4))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
            except Exception:
                pass

        p_header = doc.add_paragraph()
        p_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p_header.add_run(GJC_NAME)
        _set_font(r1, size=13, bold=True)
        p_header.add_run("\n")
        r2 = p_header.add_run(f"CUI: {GJC_CUI}  |  {GJC_ADDRESS}  |  {GJC_EMAIL}")
        _set_font(r2, size=10)
        p_header.paragraph_format.space_after = Pt(4)

    elif emitent == "candidat" and candidat_name:
        p_header = doc.add_paragraph()
        p_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p_header.add_run(candidat_name.upper())
        _set_font(r1, size=12, bold=True)
        p_header.paragraph_format.space_after = Pt(4)

    _add_horizontal_line(doc)

    # ── Dată și număr ─────────────────────────────────────────────────────────
    today = datetime.now().strftime("%d.%m.%Y")
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_date = p_date.add_run(f"Data: {today}")
    _set_font(r_date, size=11)
    p_date.paragraph_format.space_after = Pt(8)

    # ── Titlu document ────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(title.upper())
    _set_font(r_title, size=14, bold=True)
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after  = Pt(12)

    # ── Corp document ─────────────────────────────────────────────────────────
    _render_body(doc, body_text)

    # ── Bloc semnătură ────────────────────────────────────────────────────────
    doc.add_paragraph()  # spațiu
    _add_horizontal_line(doc)

    p_sign = doc.add_paragraph()
    p_sign.paragraph_format.space_before = Pt(8)
    p_sign.paragraph_format.space_after  = Pt(4)

    if emitent == "GJC":
        r_s = p_sign.add_run("Reprezentant legal GJC,\n\n\n")
        _set_font(r_s, size=11)
        r_s2 = p_sign.add_run("_______________________________\n")
        _set_font(r_s2, size=11)
        r_s3 = p_sign.add_run(f"{GJC_NAME}\nȘtampilă și semnătură")
        _set_font(r_s3, size=10)
    else:
        signer = candidat_name or "Subsemnatul/a"
        r_s = p_sign.add_run(f"{signer},\n\n\n")
        _set_font(r_s, size=11)
        r_s2 = p_sign.add_run("_______________________________\nSemnătură")
        _set_font(r_s2, size=11)

    # ── Notă audit ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p_audit = doc.add_paragraph()
    p_audit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_a = p_audit.add_run(
        f"Document generat automat de GJC AI-CRM | {today} | ID: {doc_id[:8] if doc_id else 'N/A'}"
    )
    _set_font(r_a, size=8)
    r_a.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ── Salvare ───────────────────────────────────────────────────────────────
    safe_title = re.sub(r"[^\w\-_]", "_", title)[:50]
    filename = f"{template_id}_{safe_title}_{doc_id[:8]}.docx"
    filepath = UPLOADS_DIR / filename

    doc.save(str(filepath))
    return filename


def _render_body(doc: "Document", text: str):
    """
    Transformă textul generat de Claude în paragrafe Word formatate.
    Detectează: titluri (###), liste (-), citări legale, text normal.
    """
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            # Linie goală → spațiu mic
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Titlu mare (###)
        if line.startswith("### "):
            _add_paragraph(doc, line[4:].strip(), bold=True, size=12,
                           space_before=8, space_after=4)
            i += 1
            continue

        # Titlu mediu (##)
        if line.startswith("## "):
            _add_paragraph(doc, line[3:].strip(), bold=True, size=13,
                           align=WD_ALIGN_PARAGRAPH.CENTER,
                           space_before=10, space_after=6)
            i += 1
            continue

        # Titlu mic (#)
        if line.startswith("# "):
            _add_paragraph(doc, line[2:].strip(), bold=True, size=14,
                           align=WD_ALIGN_PARAGRAPH.CENTER,
                           space_before=10, space_after=8)
            i += 1
            continue

        # Bold (**text**)
        if "**" in line:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            _render_inline_bold(p, line)
            i += 1
            continue

        # Listă (- sau •)
        if re.match(r"^[\-•]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(re.sub(r"^[\-•]\s+", "", line).strip())
            _set_font(run, size=12)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Paragraf normal
        _add_paragraph(doc, line, space_after=6)
        i += 1


def _render_inline_bold(paragraph, text: str):
    """Procesează text cu **bold** inline."""
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        _set_font(run, size=12, bold=(idx % 2 == 1))
    paragraph.paragraph_format.space_after = Pt(6)
